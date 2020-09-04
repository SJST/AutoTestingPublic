
# coding=utf-8
import os
import sys
import importlib
importlib.reload(sys)

from lmk_app.lmk_tools.resume_tools.string_process import space_process
from lmk_app.lmk_tools.resume_tools.string_process import recursive_process_UnicodeEncodeError
from lmk_app.lmk_tools.record_logs_tools.process_logs import bad_code_collection_read
from lmk_app.lmk_tools.record_logs_tools.process_logs import bad_code_clean

#读取docx中的文本代码示例
import docx
import re

'''
 解析word 文本，保存到txt文件中
'''

def read_word(resume_file_path, txt_output_path):
    #获取文档对象
    resume_file = docx.Document(resume_file_path)
    
#     paragraphs_count = len(resume_file.paragraphs)
#     tables_count = len(resume_file.tables)
#     
#     print("段落数:"+str(paragraphs_count))
#     print("表格数:"+str(tables_count))
    
    # 如果写入文件存在，则清空文件或者删除文件
    if(os.path.exists(txt_output_path)):
        os.remove(txt_output_path)
        print('exist and remove')
    
    with open(txt_output_path, 'a') as f:
        file_xml = resume_file.element.xml
        
        
        file_xml_string = str(file_xml)
        
         
        file_xml_string = xml_clean_process(file_xml_string)
        
        file_xml_string = xml_table_process(file_xml_string)
        
        
        file_xml_string = xml_paragraph_process(file_xml_string)
        
        file_xml_string = sharp_angle_process(file_xml_string)
        
#         print(file_xml_string)
        file_xml_string = text_process(f, file_xml_string)
    
#     #---    老方法，当时不会取XML    ---
#     #    docx中分表格和段落得处理
#     if tables_count > 0 and paragraphs_count < 10:
#         print('table')
#         docx_tables_process(resume_file, txt_output_path)
#     elif tables_count > 0 and paragraphs_count < 40:
#         print('both')
#         docx_tables_process(resume_file, txt_output_path)
#         docx_paragraphs_process(resume_file, txt_output_path)
#     elif tables_count == 0 and paragraphs_count > 10:
#         print('paragraph')
#         docx_paragraphs_process(resume_file, txt_output_path)
#     #---    老方法，当时不会取XML    ---

#-------------------------------------------------------------------------------------------

def xml_clean_process(file_xml_string):
    if re.search(r"<w:body>(.*?)</w:body>", file_xml_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M):
        file_xml_string = re.search(r"<w:body>(.*?)</w:body>", file_xml_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M).group(1)
    
    while(re.search(r"<w:([^>]*)Pr>(.*?)</w:\1Pr>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)):
        old_file_xml_string = re.search(r"<w:([^>]*)Pr>(.*?)</w:\1Pr>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S).group(0)
        file_xml_string = file_xml_string.replace(old_file_xml_string, "", re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
    
    while(re.search(r"<w:drawing>(.*?)</w:drawing>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)):
        old_file_xml_string = re.search(r"<w:drawing>(.*?)</w:drawing>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S).group(0)
        file_xml_string = file_xml_string.replace(old_file_xml_string, "", re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
        
    while(re.search(r"<w:hyperlink[^>]*>(.*?)</w:hyperlink>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)):
        old_file_xml_string = re.search(r"<w:hyperlink[^>]*>(.*?)</w:hyperlink>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S).group(0)
        file_xml_string = file_xml_string.replace(old_file_xml_string, "", re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
        
    while(re.search(r"<w:tblGrid[^>]*>(.*?)</w:tblGrid>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)):
        old_file_xml_string = re.search(r"<w:tblGrid[^>]*>(.*?)</w:tblGrid>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S).group(0)
        file_xml_string = file_xml_string.replace(old_file_xml_string, "", re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
    
    while(re.search(r"<mc:AlternateContent>(.*?)</mc:AlternateContent>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)):
        old_file_xml_string = re.search(r"<mc:AlternateContent>(.*?)</mc:AlternateContent>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S).group(0)
        file_xml_string = file_xml_string.replace(old_file_xml_string, "", re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
    
    #---
    if re.search(r"简历编号|最后登录", file_xml_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M):
        old_file_xml_string = re.search(r"(^.*?)<w:tbl>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S).group(1)
        file_xml_string = file_xml_string.replace(old_file_xml_string, "\n", re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
    #---
    
    while re.search(r"<w:txml:space=\"preserve\">.*?</w:t>", file_xml_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M):
        space_text_obj = re.search(r"<w:txml:space=\"preserve\">(.*?)</w:t>", file_xml_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)
        
        old_space_text_string = space_text_obj.group(0)
        new_space_text_string = space_text_obj.group(1)
        
        new_space_text_string = re.sub(r"\s+", ' ', new_space_text_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
        new_space_text_string = "@@@" + new_space_text_string + "@@@"
        
        file_xml_string = file_xml_string.replace(old_space_text_string, new_space_text_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
    
    return file_xml_string

def xml_table_process(xml_table_string):
    while(re.search(r"<w:tbl>(.*?)</w:tbl>", xml_table_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)):
        
        talbe_text_obj = re.search(r"<w:tbl>(.*?)</w:tbl>", xml_table_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)
        
        old_xml_table_string = talbe_text_obj.group(0)
        new_xml_table_string = talbe_text_obj.group(1)
        
        new_xml_table_string = re.sub(r"</w:tr>", '</w:tr>###', new_xml_table_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
        new_xml_table_string = re.sub(r"</w:tc>", '</w:tc>###', new_xml_table_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
        
        if re.search(r"tblHeader", new_xml_table_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S):
            new_xml_table_string = re.sub(r"</w:tc>", '</w:tc>:', new_xml_table_string, 1, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
        
        xml_table_string = xml_table_string.replace(old_xml_table_string, new_xml_table_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
    
    return xml_table_string

def xml_paragraph_process(xml_paragraph_string):
    while re.search("</w:p>\s*\n", xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S):
        tmp_string = re.search("</w:p>\s*\n", xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S).group(0)
        xml_paragraph_string = xml_paragraph_string.replace(tmp_string, "</w:p>###\n", re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
    
    while(re.search(r"<w:p\s*w[^>]*>(.*?)</w:p>", xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)):
        paragraph_text_obj = re.search(r"<w:p\s*w[^>]*>(.*?)</w:p>", xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)
        
        old_xml_paragraph_string = paragraph_text_obj.group(0)
        new_xml_paragraph_string = paragraph_text_obj.group(1)
        
        new_xml_paragraph_string = sharp_angle_process(new_xml_paragraph_string)
        
        new_xml_paragraph_string = re.sub(r"\s+\n\s+", '\n', new_xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
        new_xml_paragraph_string = re.sub(r"\n\s+", '\n', new_xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
        new_xml_paragraph_string = re.sub(r"([^\n]*)\n([^\n]*)", r'\1 \2', new_xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
        new_xml_paragraph_string = re.sub(r"^\s+|\s+$", '', new_xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
        
        if re.search(r"；", new_xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S):
            new_xml_paragraph_string = re.sub(r"；\s*", '：###', new_xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
        
        xml_paragraph_string = xml_paragraph_string.replace(old_xml_paragraph_string, new_xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
        xml_paragraph_string = xml_paragraph_string + '###'
#         xml_paragraph_string = re.sub(r"\n", '', xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
        
    return xml_paragraph_string

def sharp_angle_process(xml_paragraph_string):
    while(re.search(r"\s*<[^>]*>\s*", xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)):
        tmp_string = re.search(r"\s*<[^>]*>\s*", xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S).group(0)
        xml_paragraph_string = xml_paragraph_string.replace(tmp_string, "", re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
    
    while(re.search(r"#{3,}", xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S)):
        tmp_string = re.search(r"#{3,}", xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.M|re.RegexFlag.S).group(0)
        xml_paragraph_string = xml_paragraph_string.replace(tmp_string, " \n", re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M)
    
    xml_paragraph_string = re.sub(r"@{3,}", '', xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
    xml_paragraph_string = re.sub(r"\s+\n\s+", '\n', xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
    xml_paragraph_string = re.sub(r"\n\s+", '', xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
    xml_paragraph_string = re.sub(r"^\s+|\s+$", '', xml_paragraph_string, re.RegexFlag.I|re.RegexFlag.S|re.RegexFlag.M|re.RegexFlag.U)
    
    return xml_paragraph_string

def text_process(f, text):
    
    all_lines_string = bad_code_collection_read()
    text = bad_code_clean(all_lines_string, text)
    
    text = space_process(text)
    
    text = recursive_process_UnicodeEncodeError(f, text)
    
    return text

#--------------------------------------------------------------------------------------------------

# def docx_tables_process(resume_file, txt_output_path):
#     for table in resume_file.tables:
#         row_count = len(table.rows)
#         col_count = len(table.columns)
#         
#         print(str(row_count) + ' 行 ' + str(col_count) + " 列\n")
#         
#         for row in table.rows:
#             
#             duplicate_text_list = []
#             for j in range(col_count):
#                 cell = row.cells[j].text
#                 
#                 if cell in duplicate_text_list:
#                     continue
#                 else:
#                     duplicate_text_list.append(cell)
#                 
#                 with open(txt_output_path, 'a') as f:
#                     text = text_process(f, cell)
#                     print(text)
#     
# def docx_paragraphs_process(resume_file, txt_output_path):
#     #输出每一段的内容
#     for para in resume_file.paragraphs:
#         para_text = para.text
#         
#         with open(txt_output_path, 'a') as f:
#             text = text_process(f, para_text)
#             print(text)

# def read_word(resume_file_path, txt_output_path):
#     #获取文档对象
#     resume_file = docx.Document(resume_file_path)
#     print("段落数:"+str(len(resume_file.paragraphs)))
#     
#     # 如果写入文件存在，则清空文件或者删除文件
#     if(os.path.exists(txt_output_path)):
#         os.remove(txt_output_path)
#         print('exist and remove')
#         
#     #输出每一段的内容
#     for para in resume_file.paragraphs:
# #         print(para.text)
#         
#         with open(txt_output_path, 'a') as f:  #    , encoding='utf-8'
#             results = para.text
#             
#             all_lines_string = bad_code_collection_read()
#             results = bad_code_clean(all_lines_string, results)
#             
#             results = space_process(results)
#             
#             results = recursive_process_UnicodeEncodeError(f, results)
#             
#             print(results)
    
#     #输出段落编号及段落内容
#     for i in range(len(resume_file.paragraphs)):
#         print("第"+str(i)+"段的内容是："+resume_file.paragraphs[i].text)
    
if __name__ == '__main__':
    resume_file_path = "F:\\ZZMK\\ZZMK_LMK_resume_files\\all_resume\\docx\\方雅刚--众智米奇-20181031.docx"
    txt_output_path = "D:\ZZMK_LMK_resume_files\方雅刚--众智米奇-20181031.txt"
    
    read_word(resume_file_path, txt_output_path)
    
